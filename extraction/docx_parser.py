from docx import Document
import os

def extract_text_from_docx(file_path):

    pages = []

    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return []

    try:
        doc = Document(file_path)
    except Exception as e:
        print("Failed to open DOCX:", e)
        return []

    # Split document into sections by headings
    current_heading = "General"
    current_parts = []
    section_number = 1

    for para in doc.paragraphs:
        line = para.text.strip()
        if not line:
            continue

        style = para.style.name if para.style else ""

        # Detect heading styles
        is_heading = (
            "heading" in style.lower() or
            style.lower() == "title" or
            style.lower() == "subtitle"
        )

        if is_heading:
            # Save previous section
            if current_parts:
                section_text = " ".join(current_parts)
                section_text = " ".join(section_text.split())
                pages.append({
                    "page": section_number,
                    "heading": current_heading,
                    "text": section_text
                })
                section_number += 1

            current_heading = line
            current_parts = []

        else:
            current_parts.append(line)

    # Save last section
    if current_parts:
        section_text = " ".join(current_parts)
        section_text = " ".join(section_text.split())
        if section_text:
            pages.append({
                "page": section_number,
                "heading": current_heading,
                "text": section_text
            })

    # Extract tables
    for table in doc.tables:
        table_parts = []
        for row in table.rows:
            seen = set()
            for cell in row.cells:
                if cell.text not in seen:
                    seen.add(cell.text)
                    cell_text = cell.text.strip()
                    if cell_text:
                        table_parts.append(cell_text)
        if table_parts:
            pages.append({
                "page": section_number,
                "heading": current_heading,
                "text": " ".join(table_parts)
            })
            section_number += 1

    if not pages:
        print("Warning: No text extracted from DOCX.")

    return pages


if __name__ == "__main__":

    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    sample_file = os.path.join(base_dir, "data", "uploads", "06_LibreOffice_Style.docx")

    pages = extract_text_from_docx(sample_file)

    print(f"Total sections extracted: {len(pages)}")
    for p in pages:
        print(f"Section {p['page']} | Heading: {p['heading']} | {len(p['text'].split())} words")